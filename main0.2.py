import tkinter as tk
from tkinter import messagebox
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from bs4 import BeautifulSoup
from docx import Document
import time


# Função para extrair elementos com base em seus critérios
def extrair_elementos(url, tags, indices):
    chrome_options = Options()
    chrome_options.add_argument("--headless")  # Executar em segundo plano
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")

    service = Service('Users\User\Documents\chrome-headless-shell-win64\chrome-headless-shell-win64')  # Substitua pelo caminho do ChromeDriver
    driver = webdriver.Chrome(service=service, options=chrome_options)

    try:
        driver.get(url)
        time.sleep(3)  # Ajuste conforme necessário
        html = driver.page_source
    finally:
        driver.quit()  # Fecha o navegador

    soup = BeautifulSoup(html, 'html.parser')
    resultados = {}

    for tag in tags:
        elementos = soup.find_all(tag)
        resultados[tag] = []
        for i in indices.get(tag, []):
            if i < len(elementos):
                resultados[tag].append(elementos[i].get_text())
            else:
                print(f"Índice {i} fora do alcance para a tag <{tag}>.")

    return resultados

# Função para salvar os resultados em um arquivo .docx
def salvar_docx(resultados, arquivo):
    doc = Document()
    for tag, textos in resultados.items():
        doc.add_heading(f'Elementos <{tag}>', level=1)
        for texto in textos:
            doc.add_paragraph(texto)
    doc.save(arquivo)
    return arquivo

# Função chamada ao clicar no botão "Extrair"
def extrair():
    url = url_entry.get()  # Captura a URL digitada no campo
    tags = tags_entry.get().split(',')
    indices = {}
    
    # Extrair os índices das entradas
    for tag in tags:
        indices[tag.strip()] = list(map(int, indices_entry.get(tag.strip(), "").split(',')))

    try:
        resultados = extrair_elementos(url, tags, indices)
        arquivo = salvar_docx(resultados, 'resultados.docx')
        messagebox.showinfo("Sucesso", f'Resultados salvos em {arquivo}')
    except Exception as e:
        messagebox.showerror("Erro", str(e))

# Criação da janela principal
# Layout
app = tk.Tk()
app.title("Pesal - Pesquisa Salarial")

# Campo de entrada para as URLs (múltiplas)
label_urls = tk.Label(app, text="Insira as URLs (separadas por vírgula ou linha nova):")
label_urls.pack(pady=10)

entry_urls = tk.Text(app, height=5, width=60)
entry_urls.pack(pady=5)

# Campo de entrada para CBO
label_cbo = tk.Label(app, text="Insira o CBO (opcional):")
label_cbo.pack(pady=10)

entry_cbo = tk.Entry(app, width=50)
entry_cbo.pack(pady=5)

# Variável para o botão de opção
var_opcao = tk.StringVar(value="")  # Inicializa como vazia, então nenhum botão estará selecionado

# Botões de opção para "Salário" e "Dissídio"
radio_salario = tk.Radiobutton(app, text="Salário", variable=var_opcao, value="salario")
radio_salario.pack(pady=5)

radio_dissidio = tk.Radiobutton(app, text="Dissídio", variable=var_opcao, value="dissidio")
radio_dissidio.pack(pady=5)

# Botão para coletar parágrafos

tk.Label(app, text="Tags (ex: p, h1, h2):").grid(row=1, column=0, padx=10, pady=5)
tags_entry = tk.Entry(app, width=50)
tags_entry.grid(row=1, column=1, padx=10, pady=5)

tk.Label(app, text="Índices (ex: 0, 1):").grid(row=2, column=0, padx=10, pady=5)
indices_entry = tk.Entry(app, width=50)
indices_entry.grid(row=2, column=1, padx=10, pady=5)

extrair_button = tk.Button(app, text="Extrair", command=extrair)
extrair_button.grid(row=3, column=0, columnspan=2, pady=20)

# Iniciar a GUI
app.mainloop()



