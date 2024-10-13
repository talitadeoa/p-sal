import tkinter as tk
from tkinter import messagebox
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from bs4 import BeautifulSoup
from docx import Document
import time
from webdriver_manager.chrome import ChromeDriverManager

# Função para extrair elementos com base em critérios predefinidos
def extrair_elementos(url):
    # Opções do Chrome para execução em segundo plano
    chrome_options = Options()
    chrome_options.add_argument("--headless")  # Executar em segundo plano
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")

    service = Service(ChromeDriverManager().install())
    driver = webdriver.Chrome(service=service, options=chrome_options)

    try:
        driver.get(url)
        time.sleep(3)  # Espera para garantir que a página carregue (ajuste conforme necessário)
        html = driver.page_source
    finally:
        driver.quit()  # Fecha o navegador

    soup = BeautifulSoup(html, 'html.parser')
    resultados = {}

    # Tags e índices predefinidos
    tags = {
        'p': [0, 1],  # Extrai os dois primeiros parágrafos
        'h1': [0],    # Extrai o primeiro título h1
        'h2': [0, 1]  # Extrai os dois primeiros títulos h2
    }

    for tag, indices in tags.items():
        elementos = soup.find_all(tag)
        resultados[tag] = []
        for i in indices:
            if i < len(elementos):
                resultados[tag].append(elementos[i].get_text())
            else:
                print(f"Índice {i} fora do alcance para a tag <{tag}>.")

    return resultados

# Função para salvar os resultados em um arquivo .docx
def salvar_docx(resultados, arquivo):
    doc = Document()
    for tag, textos in resultados.items():
        doc.add_heading(f'Elementos <{tag}>', level=1)
        for texto in textos:
            doc.add_paragraph(texto)
    doc.save(arquivo)
    return arquivo

# Função chamada ao clicar no botão "Extrair"
def extrair():
    url = url_entry.get()  # Captura a URL digitada no campo
    try:
        resultados = extrair_elementos(url)
        arquivo = salvar_docx(resultados, 'resultados.docx')
        messagebox.showinfo("Sucesso", f'Resultados salvos em {arquivo}')
    except Exception as e:
        messagebox.showerror("Erro", str(e))

# Criação da janela principal
root = tk.Tk()
root.title("Web Scraper")

# Layout
tk.Label(root, text="URL:").grid(row=0, column=0, padx=10, pady=5)
url_entry = tk.Entry(root, width=50)  # Campo de entrada para a URL
url_entry.grid(row=0, column=1, padx=10, pady=5)

extrair_button = tk.Button(root, text="Extrair", command=extrair)
extrair_button.grid(row=1, column=0, columnspan=2, pady=20)

root.mainloop()
