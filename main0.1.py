import tkinter as tk
from tkinter import messagebox
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from bs4 import BeautifulSoup
import docx
from concurrent.futures import ThreadPoolExecutor
from webdriver_manager.chrome import ChromeDriverManager

# Defina as opções do Chrome
chrome_options = Options()
chrome_options.add_argument("--headless")  # Executa o Chrome em segundo plano

# Função para coletar parágrafos de uma URL
def coletar_paragrafos(url, paragrafo_inicial):
    try:
        # Usando Service e ChromeDriverManager para gerenciar o ChromeDriver
        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=chrome_options)

        # Acessa a página
        driver.get(url)
        soup = BeautifulSoup(driver.page_source, 'html.parser')
        driver.quit()

        # Filtra parágrafos que começam com o texto específico
        parrafos = soup.find_all('p')
        resultado = [p.get_text() for p in parrafos if p.get_text().startswith(paragrafo_inicial)]
        return url, resultado
    except Exception as e:
        return url, f"Erro ao processar a URL {url}: {str(e)}"

# Função para salvar parágrafos em um arquivo DOCX
def salvar_em_docx(paragrafos_por_url, nome_arquivo):
    doc = docx.Document()
    for url, paragrafos in paragrafos_por_url.items():
        doc.add_heading(f'Parágrafos de {url}', level=1)
        if isinstance(paragrafos, list):
            for paragrafo in paragrafos:
                doc.add_paragraph(paragrafo)
        else:
            doc.add_paragraph(paragrafos)  # Adiciona mensagem de erro, se houver
    doc.save(nome_arquivo)

# Função para executar a coleta de URLs em paralelo
def processar_urls_em_paralelo(urls, paragrafo_inicial):
    with ThreadPoolExecutor() as executor:
        results = executor.map(lambda url: coletar_paragrafos(url, paragrafo_inicial), urls)
    return dict(results)

# Função de ação do botão
def acao_coletar():
    urls = entry_urls.get("1.0", "end-1c").split(",")  # Divide as URLs por vírgula
    cbo = entry_cbo.get()
    selecionado = var_opcao.get()

    if cbo:  # Verifica se há algo no campo CBO
        messagebox.showinfo("Atenção", "Ainda não é possível pesquisar por CBO, mas você pode realizar a pesquisa através da URL.")
        return

    if not urls or all(url.strip() == "" for url in urls):
        messagebox.showerror("Erro", "Por favor, insira pelo menos uma URL.")
        return
    
    if not selecionado:
        messagebox.showerror("Erro", "Por favor, selecione uma opção.")
        return

    # Remove URLs vazias e espaços em branco
    urls = [url.strip() for url in urls if url.strip()]

    # Define os parágrafos iniciais com base na escolha do usuário
    if selecionado == "Salario":
        paragrafo_inicial = 'Seu trecho inicial para Salário'  # Defina o texto inicial específico para Salário
    elif selecionado == "Dissidio":
        paragrafo_inicial = 'Seu trecho inicial para Dissídio'  # Defina o texto inicial específico para Dissídio

    # Executa a coleta das URLs em paralelo com o parágrafo inicial selecionado
    paragrafos_por_url = processar_urls_em_paralelo(urls, paragrafo_inicial)

    if not paragrafos_por_url:
        messagebox.showinfo("Resultado", "Nenhum parágrafo encontrado.")
        return
    
    salvar_em_docx(paragrafos_por_url, 'pesquisa_salarial.docx')
    messagebox.showinfo("Sucesso", "Dados salvos em 'pesquisa_salarial.docx'.")

# Configuração da GUI
app = tk.Tk()
app.title("Pesal - Pesquisa Salarial")
app.geometry("500x400")

# Campo de entrada para as URLs (múltiplas)
label_urls = tk.Label(app, text="Insira as URLs (separadas por vírgula ou linha nova):")
label_urls.pack(pady=10)

entry_urls = tk.Text(app, height=5, width=60)
entry_urls.pack(pady=5)

# Campo de entrada para CBO
label_cbo = tk.Label(app, text="Insira o CBO (opcional):")
label_cbo.pack(pady=10)

entry_cbo = tk.Entry(app, width=50)
entry_cbo.pack(pady=5)

# Variável para o botão de opção
var_opcao = tk.StringVar(value="")  # Inicializa como vazia, então nenhum botão estará selecionado

# Botões de opção para "Salário" e "Dissídio"
radio_salario = tk.Radiobutton(app, text="Salário", variable=var_opcao, value="salario")
radio_salario.pack(pady=5)

radio_dissidio = tk.Radiobutton(app, text="Dissídio", variable=var_opcao, value="dissidio")
radio_dissidio.pack(pady=5)

# Botão para coletar parágrafos
botao_coletar = tk.Button(app, text="Coletar Dados", command=acao_coletar)
botao_coletar.pack(pady=20)

# Iniciar a GUI
app.mainloop()
