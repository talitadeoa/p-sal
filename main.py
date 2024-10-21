import tkinter as tk
from tkinter import messagebox
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from bs4 import BeautifulSoup
from docx import Document
import time
import threading
import re
from webdriver_manager.chrome import ChromeDriverManager

# Função para extrair elementos com base em critérios predefinidos
def extrair_elementos(url, tipo):
    # Opções do Chrome para execução em segundo plano
    chrome_options = Options()
    chrome_options.add_argument("--headless")  # Executar em segundo plano
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")

    service = Service(ChromeDriverManager().install())
    driver = webdriver.Chrome(service=service, options=chrome_options)

    try:
        driver.get(url)
        time.sleep(3)  # Espera para garantir que a página carregue (ajuste conforme necessário)
        html = driver.page_source
    finally:
        driver.quit()  # Fecha o navegador

    soup = BeautifulSoup(html, 'html.parser')
    resultados = {}

    if tipo == "salario":
        # Expressão regular para capturar as partes da URL
        padrao = r'https://www\.salario\.com\.br/profissao/([^/]+)-cbo-([0-9]+)/?(.*)'
        match = re.match(padrao, url)

        if match:
            cargo = match.group(1).replace('-', ' ')
            cbo = match.group(2)
            local_info = match.group(3)  # Captura tudo após o CBO, se houver

            # Determinar o tipo de local
            if not local_info:  # Caso 1: apenas a estrutura básica (país)
                local = ('país', 'br')
            elif len(local_info) == 2:  # Caso 2: estado
                local = ('estado', local_info)
            else:  # Caso 3: cidade
                local = ('cidade', local_info.strip('/'))

            # Tags e índices predefinidos para o tipo "salario"
            tags = {
                'p': [0, 1, 2, 7],
                'h2': [1, 13],
                'table': [3],
            }

    elif tipo == "dissidio":
        cargo = url.split('/')[-2].replace('-', ' ').capitalize()
        cbo = soup.find_all('a')[3].get_text().strip()  # Captura o cargo

        tags = {
            'p': [24, 26, 27, 28, 29],
            'h2': [10, 11]
        }

    for tag, indices in tags.items():
        elementos = soup.find_all(tag)
        resultados[tag] = []
        for i in indices:
            if i < len(elementos):
                resultados[tag].append(elementos[i].get_text())
            else:
                print(f"Índice {i} fora do alcance para a tag <{tag}>.")

    # Salva os resultados no documento específico
    if tipo == "salario":
        arquivo = f'Salário {cargo} - CBO {cbo} - {local[1]}.docx'
    elif tipo == "dissidio":
        arquivo = f'Dissídio {cargo} - {cbo}.docx'
    salvar_docx(resultados, arquivo)

# Função para salvar os resultados em um arquivo .docx
def salvar_docx(resultados, arquivo):
    doc = Document()
    for tag, textos in resultados.items():
        doc.add_heading(f'Elementos <{tag}>', level=1)
        for texto in textos:
            doc.add_paragraph(texto)
    doc.save(arquivo)

# Função chamada ao clicar no botão "Extrair"
def extrair():
    urls = url_entry.get().split(',')  # Captura as URLs digitadas no campo, separando por vírgula
    tipo_selecionado = tipo_var.get()

    if not urls:
        messagebox.showerror("Erro", "Por favor, insira as URLs.")
        return

    threads = []
    for url in urls:
        url = url.strip()
        thread = threading.Thread(target=extrair_elementos, args=(url, tipo_selecionado))
        threads.append(thread)
        thread.start()

    for thread in threads:
        thread.join()

    messagebox.showinfo("Sucesso", "Dados coletados e salvos com sucesso!")

# Criação da janela principal
root = tk.Tk()
root.title("Automação - Pesquisa Salarial")
root.geometry("500x400")

# Variável para armazenar a opção selecionada pelo usuário
tipo_var = tk.StringVar(value="salario")  # Valor padrão

# Layout
tk.Label(root, text="Insira as URLs (separadas por vírgula):").grid(row=0, column=0, padx=10, pady=10)
url_entry = tk.Entry(root, width=50)  # Campo de entrada para a URL
url_entry.grid(row=0, column=1, padx=10, pady=5)

# Opções de seleção para o tipo de extração
tk.Radiobutton(root, text="Salário", variable=tipo_var, value="salario").grid(row=3, column=0, padx=10, pady=5)
tk.Radiobutton(root, text="Dissídio", variable=tipo_var, value="dissidio").grid(row=3, column=1, padx=10, pady=5)

# Botão para iniciar a extração
extrair_button = tk.Button(root, text="Extrair", command=extrair)
extrair_button.grid(row=4, column=0, columnspan=2, pady=20)

root.mainloop()
